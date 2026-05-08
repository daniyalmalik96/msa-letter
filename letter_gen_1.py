# ============================================================
# MSA Letter Generator
# Install requirements: pip install streamlit python-docx
# Run: streamlit run letter_gen.py
# ============================================================

import io
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ═══════════════════════════════════════════════════════════════
#  HELPERS
# ═══════════════════════════════════════════════════════════════

FONT = "Arial"


def _set_font(run, size_pt: float, bold=False, italic=False):
    run.font.name   = FONT
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), FONT)
    ex = rPr.find(qn("w:rFonts"))
    if ex is not None:
        rPr.remove(ex)
    rPr.insert(0, rFonts)


def _set_para_fmt(para, before_pt=0, after_pt=6):
    fmt = para.paragraph_format
    fmt.space_before      = Pt(before_pt)
    fmt.space_after       = Pt(after_pt)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _blank(doc, after_pt=6):
    p = doc.add_paragraph()
    _set_para_fmt(p, before_pt=0, after_pt=after_pt)
    return p


def _para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT,
          size_pt=12, bold=False, before=0, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    _set_para_fmt(p, before_pt=before, after_pt=after)
    if text:
        r = p.add_run(text)
        _set_font(r, size_pt, bold=bold)
    return p


def _set_indent_xml(para, left_twips=0, first_line_twips=0, hanging_twips=0):
    """
    Set paragraph indentation directly via XML.
    left_twips       — left indent from margin
    first_line_twips — extra first-line indent (positive)
    hanging_twips    — hanging indent (wrapping lines pushed right by this)
    """
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left_twips))
    if hanging_twips:
        ind.set(qn("w:hanging"), str(hanging_twips))
    elif first_line_twips:
        ind.set(qn("w:firstLine"), str(first_line_twips))
    ex = pPr.find(qn("w:ind"))
    if ex is not None:
        pPr.remove(ex)
    pPr.append(ind)


def _set_tab_stops(para, positions):
    """Set left-aligned tab stops (twips) on a paragraph."""
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    for pos in positions:
        t = OxmlElement("w:tab")
        t.set(qn("w:val"), "left")
        t.set(qn("w:pos"), str(pos))
        tabs.append(t)
    ex = pPr.find(qn("w:tabs"))
    if ex is not None:
        pPr.remove(ex)
    pPr.append(tabs)


# ═══════════════════════════════════════════════════════════════
#  HEADER  (true Word header zone)
# ═══════════════════════════════════════════════════════════════

def _build_header(section):
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    def hpara(text, size_pt=12, bold=False, after=0):
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_fmt(p, before_pt=0, after_pt=after)
        if text:
            r = p.add_run(text)
            _set_font(r, size_pt, bold=bold)

    hpara("CONFIDENTIAL",             size_pt=12, bold=True,  after=6)
    hpara("",                                                  after=0)
    hpara("PAF Hospital, Islamabad",  size_pt=14, bold=True,  after=0)
    hpara("(Medical Staff Affairs)",  size_pt=12, bold=False, after=0)


# ═══════════════════════════════════════════════════════════════
#  FOOTER  (true Word footer zone)
# ═══════════════════════════════════════════════════════════════

def _build_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_fmt(p, before_pt=0, after_pt=0)
    r = p.add_run("CONFIDENTIAL")
    _set_font(r, 12, bold=True)


# ═══════════════════════════════════════════════════════════════
#  NUMBERED LIST ITEM
#
#  Rules:
#    • Left indent  = 0  (text starts exactly at the left margin)
#    • First-line indent = 0  (number also at the left margin)
#    • Tab stop at 432 twips (≈ 3 char widths / 0.3") so text
#      after the number is consistently indented a short gap
#    • Wrapped lines align with the text start (hanging = 432)
# ═══════════════════════════════════════════════════════════════

TAB_NUM = 432   # ~0.3 inch — gap between number and text


def _numbered_item(doc, number: int, text: str, size_pt=12, after_pt=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_fmt(p, before_pt=0, after_pt=after_pt)

    # No left indent from margin; hanging keeps wrapped lines at text start
    _set_indent_xml(p, left_twips=TAB_NUM, hanging_twips=TAB_NUM)
    _set_tab_stops(p, [TAB_NUM])

    r = p.add_run(f"{number}.\t{text}")
    _set_font(r, size_pt, bold=False)
    return p


# ═══════════════════════════════════════════════════════════════
#  DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════

def generate_document(recipient: str, subject: str,
                      first_line: str, second_line: str,
                      dated: str) -> bytes:

    doc = Document()

    # Default Normal style
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    nfmt = normal.paragraph_format
    nfmt.space_before      = Pt(0)
    nfmt.space_after       = Pt(6)
    nfmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # ── Page: A4, margins ─────────────────────────────────────
    section = doc.sections[0]
    section.page_width      = Cm(21.0)
    section.page_height     = Cm(29.7)
    section.left_margin     = Inches(1.5)   # 1.5" left
    section.right_margin    = Inches(0.5)   # 0.5" right
    section.top_margin      = Inches(0.5)   # 0.5" top
    section.bottom_margin   = Inches(0.5)   # 0.5" bottom
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    _build_header(section)
    _build_footer(section)

    # ══════════════════════════════════════════════════════════
    #  RECIPIENT — left, bold
    # ══════════════════════════════════════════════════════════
    _para(doc, recipient,
          align=WD_ALIGN_PARAGRAPH.LEFT,
          size_pt=12, bold=True, before=0, after=12)

    # ══════════════════════════════════════════════════════════
    #  SUBJECT — CENTER aligned, bold, 16pt
    # ══════════════════════════════════════════════════════════
    _para(doc, subject,
          align=WD_ALIGN_PARAGRAPH.CENTER,
          size_pt=16, bold=True, before=0, after=12)

    # ══════════════════════════════════════════════════════════
    #  BODY — numbered, at left margin, blank line between items
    # ══════════════════════════════════════════════════════════
    _numbered_item(doc, 1, first_line, after_pt=0)
    _blank(doc, after_pt=6)

    _numbered_item(doc, 2, second_line, after_pt=0)
    _blank(doc, after_pt=6)

    _numbered_item(doc, 3, "Thankyou.", after_pt=0)
    _blank(doc, after_pt=14)

    # ══════════════════════════════════════════════════════════
    #  SIGNATURE BLOCK
    #
    #  All four lines share a fixed left indent of 3.5" (5040 twips)
    #  from the page margin.  Because left_twips is set (not a tab),
    #  ANY wrapped line — including "Alauddin Ahmed & Cecil Chaudhry
    #  Complex" — automatically continues below "Alauddin", never
    #  falling back to the extreme left.
    # ══════════════════════════════════════════════════════════

    SIG_LEFT = 5040   # 3.5" from left margin → right half of page

    sig_lines = [
        ("(DR AMIR HALEEM)",                        True),
        ("Medical Director",                        False),
        ("PAF Hospital, Islamabad",                 False),
        ("Alauddin Ahmed & Cecil Chaudhry Complex", False),
    ]

    for line_text, is_bold in sig_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_para_fmt(p, before_pt=0, after_pt=2)
        # Left indent at 3.5" — wrapping stays at the same edge
        _set_indent_xml(p, left_twips=SIG_LEFT)
        r = p.add_run(line_text)
        _set_font(r, 12, bold=is_bold)

    _blank(doc, after_pt=10)

    # ══════════════════════════════════════════════════════════
    #  LM NUMBER — left, regular
    # ══════════════════════════════════════════════════════════
    _para(doc, f"LM No. IH/79101/1/MD dated: {dated}",
          align=WD_ALIGN_PARAGRAPH.LEFT,
          size_pt=12, bold=False, before=0, after=12)

    # ══════════════════════════════════════════════════════════
    #  COPY TO — bold label + numbered entries at left margin
    #
    #  Column layout (tab stops):
    #    number  → left margin (0)
    #    office  → TAB_NUM (432 twips, ~0.3")
    #    colon   → 3600 twips (2.5")
    #    action  → 4320 twips (3.0")
    # ══════════════════════════════════════════════════════════
    _para(doc, "Copy to:",
          align=WD_ALIGN_PARAGRAPH.LEFT,
          size_pt=12, bold=True, before=0, after=4)

    TAB_COLON  = 3600
    TAB_ACTION = 4320

    copy_entries = [
        ("SO to DG(MS)",        "For Info, Please"),
        ("Director Operations", "'-'"),
    ]

    for idx, (office, action) in enumerate(copy_entries, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_para_fmt(p, before_pt=0, after_pt=4)
        # Same indent rules as body: starts at left margin, text after tab
        _set_indent_xml(p, left_twips=TAB_NUM, hanging_twips=TAB_NUM)
        _set_tab_stops(p, [TAB_NUM, TAB_COLON, TAB_ACTION])
        r = p.add_run(f"{idx}.\t{office}\t:\t{action}")
        _set_font(r, 12, bold=False)

    # ── Serialize ─────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════
#  STREAMLIT UI
# ═══════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="MSA Letter Generator",
    page_icon="📄",
    layout="centered",
)

st.markdown("""
<style>
    .block-container { padding-top: 2rem; max-width: 820px; }
    .stTextInput > label, .stTextArea > label {
        font-weight: 600; font-size: 0.9rem;
    }
    .stButton > button { border-radius: 6px; }
</style>
""", unsafe_allow_html=True)

st.markdown("## 📄  MSA Letter Generator")
st.caption(
    "Fill in the five fields and click **Generate Letter** "
    "to download a fully formatted `.docx` file."
)
st.divider()

col1, col2 = st.columns([3, 2])
with col1:
    recipient = st.text_input(
        "Recipient",
        placeholder="e.g.  TO WHOM IT MAY CONCERN",
        help="Appears bold, left-aligned.",
    )
    subject = st.text_input(
        "Subject",
        placeholder="e.g.  POSTING OF MEDICAL OFFICER",
        help="Centre-aligned, bold, 16 pt.",
    )
with col2:
    dated = st.text_input(
        "Date  (DD Mon YYYY)",
        placeholder="e.g.  15 Mar 2025",
        help="Goes into the LM No. reference line.",
    )

first_line = st.text_area(
    "Body Point 1",
    placeholder="First numbered point of the letter…",
    height=90,
)
second_line = st.text_area(
    "Body Point 2",
    placeholder="Second numbered point of the letter…",
    height=90,
)

st.divider()

if st.button("⚡  Generate Letter", type="primary", use_container_width=True):
    checks = {
        "Recipient":    recipient,
        "Subject":      subject,
        "Date":         dated,
        "Body Point 1": first_line,
        "Body Point 2": second_line,
    }
    missing = [k for k, v in checks.items() if not v.strip()]
    if missing:
        st.warning(f"Please fill in: **{', '.join(missing)}**")
    else:
        with st.spinner("Building document…"):
            docx_bytes = generate_document(
                recipient   = recipient.strip(),
                subject     = subject.strip(),
                first_line  = first_line.strip(),
                second_line = second_line.strip(),
                dated       = dated.strip(),
            )
        st.success("✅  Letter ready — click below to download.")
        st.download_button(
            label     = "⬇️  Download Letter (.docx)",
            data      = docx_bytes,
            file_name = "MSA_Letter.docx",
            mime      = (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            use_container_width=True,
        )

with st.expander("📋  Format Specification Reference"):
    st.markdown("""
| Element | Specification |
|---|---|
| Paper | A4 Portrait |
| Font | Arial, 12 pt body · 16 pt subject |
| Margins | Left 1.5" · Right / Top / Bottom 0.5" |
| Header (Word zone) | CONFIDENTIAL · blank · Hospital name · Dept |
| Footer (Word zone) | CONFIDENTIAL |
| Recipient | Left-aligned, bold, 12 pt |
| Subject | **Centre-aligned**, bold, 16 pt |
| Body points | Numbered 1. 2. 3. — start at left margin, text 0.3" after number, wraps aligned |
| Closing item | Fixed: *3. Thankyou.* |
| Signature | Left edge fixed at 3.5" — wrapping stays below "Alauddin", never falls left |
| LM No. | Left-aligned, regular |
| Copy to | Numbered 1. 2. — same left-margin rules, colon-aligned columns |
""")
